import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import os
from dotenv import load_dotenv
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import logging
from docx import Document
from docx2pdf import convert
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
EMAIL_SENDER = os.getenv("EMAIL_SENDER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
SMTP_SERVER = os.getenv("SMTP_SERVER")
SMTP_PORT = int(os.getenv("SMTP_PORT"))

# Directory for certificates
CERTIFICATE_DIR = "certificates"
os.makedirs(CERTIFICATE_DIR, exist_ok=True)

# Word template path
TEMPLATE_PATH = r"D:\ecertificate\templates\Styled_Certificate - Copy.docx"

def add_page_border_fixed(section):
    pgBorders = OxmlElement('w:pgBorders')
    for border_type in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # 24 = 3pt
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'D4AF37')  # Gold color
        pgBorders.append(border)
    section._sectPr.insert(2, pgBorders)

def send_email(to_email, name, participation_item, attachment_path):
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_SENDER
        msg['To'] = to_email
        msg['Subject'] = f"Certificate of Participation - {participation_item}"

        body = f"Dear {name},\n\nCongratulations on your participation in {participation_item} at  Techtrix  2025, organized by Department of Computer Science, St. Joseph’s Academy of Higher Education & Research! Please find your certificate attached.\n\nBest regards,\nDepartment of Computer Science Team"
        msg.attach(MIMEText(body, 'plain'))

        with open(attachment_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename={os.path.basename(attachment_path)}"
        )
        msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.sendmail(EMAIL_SENDER, to_email, msg.as_string())
        logger.info(f"Email sent successfully to {to_email} for {name}")
    except Exception as e:
        logger.error(f"Email sending failed for {name} to {to_email}: {str(e)}")
        raise

def generate_certificate(name, participation_item, output_path):
    try:
        logger.info(f"Generating certificate for {name}")
        
        # Load the Word template
        if not os.path.exists(TEMPLATE_PATH):
            logger.error(f"Template not found at {TEMPLATE_PATH}")
            raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")
        
        doc = Document(TEMPLATE_PATH)
        
        # Set page size and margins
        section = doc.sections[0]
        new_width_cm, new_height_cm = 27.94, 21.59  # 21.59 x 35.56 cm
        section.page_width = Cm(new_width_cm)
        section.page_height = Cm(new_height_cm)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        
        # Add border
        add_page_border_fixed(section)
        
        # Log initial page size and paragraph content
        logger.debug(f"Initial page size: {new_width_cm} x {new_height_cm} cm")
        initial_paragraphs = [(i, p.text) for i, p in enumerate(doc.paragraphs)]
        logger.debug(f"Initial paragraphs: {initial_paragraphs}")
        
        # Replace placeholders and set font sizes
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                original_text = run.text
                if '{name}' in run.text:
                    logger.debug(f"Paragraph {i}, Run: Replacing '{original_text}' with '{name}'")
                    run.text = run.text.replace('{name}', name)
                    run.font.name = 'Palatino Linotype'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Palatino Linotype')
                    run.font.size = Pt(30)
                    run.bold = True
                    run.font.color.rgb = RGBColor(12, 45, 90)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif '{participation_item}' in run.text or '(participation_item)' in run.text:
                    logger.debug(f"Paragraph {i}, Run: Replacing '{original_text}' with '{participation_item}'")
                    run.text = run.text.replace('{participation_item}', participation_item).replace('(participation_item)', participation_item)
                    run.font.name = 'Palatino Linotype'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Palatino Linotype')
                    run.font.size = Pt(22)
                    run.italic = True
                    run.bold = True  # To match potential *** or bold-italic intent
                    run.font.color.rgb = RGBColor(12, 45, 90)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif '{date}' in run.text:
                    logger.debug(f"Paragraph {i}, Run: Replacing '{original_text}' with '{datetime.now().strftime('%B %d, %Y')}'")
                    run.text = run.text.replace('{date}', datetime.now().strftime("%B %d, %Y"))
                    run.font.name = 'Palatino Linotype'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Palatino Linotype')
                    run.font.size = Pt(16)  # Adjust to Pt(30) if preferred
                    run.font.color.rgb = RGBColor(12, 45, 90)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Apply italic to lines with * markup
                elif '*' in run.text and not any(placeholder in run.text for placeholder in ['{name}', '{participation_item}', '{date}']):
                    run.italic = True
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Log updated paragraph content
        updated_paragraphs = [(i, p.text) for i, p in enumerate(doc.paragraphs)]
        logger.debug(f"Updated paragraphs: {updated_paragraphs}")
        
        # Save the modified document to a temporary file
        temp_doc_path = os.path.join(CERTIFICATE_DIR, f"{name}_temp.docx")
        doc.save(temp_doc_path)
        logger.info(f"Temporary file saved at {temp_doc_path}")
        
        # Convert to PDF
        logger.info(f"Converting to PDF for {name} at {output_path}")
        convert(temp_doc_path, output_path)
        logger.info(f"PDF generated successfully for {name}")
        
        # Clean up temporary file
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)
    except Exception as e:
        logger.error(f"PDF generation failed for {name}: {str(e)}")
        raise

def process_excel(file_path, result_text):
    try:
        # Validate file existence and readability
        if not os.path.exists(file_path):
            messagebox.showerror("Error", f"Excel file not found at {file_path}")
            logger.error(f"Excel file not found at {file_path}")
            return
        if os.path.getsize(file_path) == 0:
            messagebox.showerror("Error", f"Excel file at {file_path} is empty")
            logger.error(f"Excel file at {file_path} is empty")
            return

        # Read Excel file
        logger.info(f"Reading Excel file from {file_path}")
        df = pd.read_excel(file_path)
        logger.debug(f"Columns found: {df.columns.tolist()}")
        if not all(col in df.columns for col in ['Name', 'Participation Item', 'Email']):
            messagebox.showerror("Error", "Excel file must contain 'Name', 'Participation Item', and 'Email' columns")
            logger.error("Missing required columns in Excel file")
            return

        results = []
        for index, row in df.iterrows():
            name = row['Name']
            participation_item = row['Participation Item']
            email = row['Email']
            certificate_path = os.path.join(CERTIFICATE_DIR, f"{name}_certificate.pdf")

            # Generate certificate
            generate_certificate(name, participation_item, certificate_path)

            # Send email
            try:
                logger.info(f"Sending email to {email} for {name}")
                send_email(email, name, participation_item, certificate_path)
                results.append(f"{name} ({email}): Sent")
                os.remove(certificate_path)  # Clean up
                logger.info(f"Email and cleanup completed for {name}")
            except Exception as e:
                results.append(f"{name} ({email}): Failed - {str(e)}")
                logger.error(f"Email failed for {name}: {str(e)}")

        result_text.delete(1.0, tk.END)
        result_text.insert(tk.END, "\n".join(results))
        messagebox.showinfo("Success", "Processing complete! Check results below.")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to process Excel file: {str(e)}")
        logger.error(f"Failed to process Excel file: {str(e)}")

def create_gui():
    root = tk.Tk()
    root.title("Certificate Generator")
    root.geometry("600x400")

    file_label = tk.Label(root, text="Select Excel File:")
    file_label.pack(pady=10)

    file_path_var = tk.StringVar()
    file_entry = tk.Entry(root, textvariable=file_path_var, width=50)
    file_entry.pack(pady=5)

    def browse_file():
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        file_path_var.set(file_path)

    browse_button = tk.Button(root, text="Browse", command=browse_file)
    browse_button.pack(pady=5)

    result_label = tk.Label(root, text="Results:")
    result_label.pack(pady=10)

    result_text = tk.Text(root, height=10, width=60)
    result_text.pack(pady=5)

    def start_processing():
        file_path = file_path_var.get()
        if not file_path:
            messagebox.showerror("Error", "Please select an Excel file")
            return
        process_excel(file_path, result_text)

    process_button = tk.Button(root, text="Generate & Send Certificates", command=start_processing)
    process_button.pack(pady=10)

    root.mainloop()

if __name__ == "__main__":
    create_gui()    